"""
Word 文档生成器 - 使用 python-docx 生成文档
"""
import os
from typing import Dict, List, Any
from docx import Document
from docx.shared import Inches


class DocxWriter:
    """Word 文档生成器"""

    def __init__(self, output_path: str):
        self.output_path = output_path
        self.doc = Document()

        self.chapter_index = 0

    def add_title(self, title: str):
        self.doc.add_heading(title, level=0)

    def add_chapter(self, chapter: Dict[str, Any]):
        self.chapter_index += 1
        idx = self.chapter_index

        name = chapter.get('name', '')
        description = chapter.get('description', '')
        functions = chapter.get('functions', [])

        # H1: 章节标题
        self.doc.add_heading(f'{idx}. {name}', level=1)

        # H2: 1. 产品概述
        self.doc.add_heading(f'{idx}.1. 产品概述', level=2)
        self.doc.add_paragraph(description)

        # H2: 2. 产品结构（功能摘要)
        self.doc.add_heading(f'{idx}.2. 产品结构（功能摘要)', level=2)
        self.doc.add_paragraph('产品结构如图：')

        structure_image = chapter.get('structure_image', '')
        if structure_image and os.path.exists(structure_image):
            try:
                self.doc.add_picture(structure_image, width=Inches(7))
            except Exception:
                self.doc.add_paragraph('结构图生成失败，未能插入图片。')
        else:
            self.doc.add_paragraph('结构图生成失败，未能插入图片。')

        self.doc.add_paragraph('主要包括如下功能')
        for fn in functions:
            self.doc.add_paragraph(fn, style='List Bullet')

        # H2: 3. 特性说明
        self.doc.add_heading(f'{idx}.3. 特性说明', level=2)

        for loc, feature in enumerate(chapter.get('features', []), start=1):
            scenario = feature.get('scenario', '')
            process = feature.get('process', [])
            input_data = feature.get('input', '')
            output_data = feature.get('output', '')
            flow_chart = feature.get('flow_chart', '')

            # H3: 特性子标题 (用户场景)
            self.doc.add_heading(f'{idx}.3.{loc}. {scenario}', level=3)
            self.doc.add_paragraph(f'用户场景： {scenario}', style='List Bullet')

            self.doc.add_paragraph('流程图', style='List Bullet')
            if flow_chart and os.path.exists(flow_chart):
                try:
                    self.doc.add_picture(flow_chart, width=Inches(7))
                except Exception:
                    self.doc.add_paragraph('流程图生成失败，未能插入图片。')
            else:
                self.doc.add_paragraph('流程图生成失败，未能插入图片。')

            if process:
                self.doc.add_paragraph('功能过程:', style='List Bullet')
                for step, order in enumerate(process, start=1):
                    self.doc.add_paragraph(f'{step}. {order}'.rstrip())

            self.doc.add_paragraph(f'输入：{input_data}', style='List Bullet')
            self.doc.add_paragraph(f'输出：{output_data}', style='List Bullet')

    def save(self) -> str:
        self.doc.save(self.output_path)
        return self.output_path
