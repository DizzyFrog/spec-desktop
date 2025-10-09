import json
from collections import OrderedDict
import pandas as pd
from typing import Optional
import os
from pathlib import Path
from .myconfig import config_manager
from app.log import logger

from .content import Chapter
from .feature import Feature
from .tool import get_description, get_structure, get_flow_chart, get_description_async

from docx import Document
from docx.shared import Inches
import concurrent.futures
from functools import partial

class MetaDataProcessor:
    """处理Excel元数据的类"""
    
    def __init__(self, handle_duplicates="merge"):
        """初始化MetaDataProcessor，从配置文件读取参数
        
        Args:
            handle_duplicates: 处理重复功能用户需求的方式
                - "merge": 合并到同一个需求下（默认）
                - "separate": 创建新的版本（如：需求_v2）
                - "error": 遇到重复时报错
        """
        config = config_manager.config
        self.file_name = Path("web/public/resource/file/uploads") / config["excel_file_name"]
        self.sheet_name = config["sheet_name"]
        self.df: Optional[pd.DataFrame] = None
        # self.columns = ["B", "C", "D", "E", "F", "H", "L"]
        self.columns = ['功能用户需求', '触发事件', '功能过程', '子过程描述', '数据组', '功能用户', '角色']
        self.output_path = Path("web/public/resource/file/output/data.json")
        self.handle_duplicates = handle_duplicates
        
        self.progress_file = Path("web/public/resource/file/output/progress.json")
        self.progress = {
            "current_key": "",
            "total": 0,
            "current": 0,
            "percentage": 0,
            "status": "idle"  # idle, processing, completed, error
        }
        # 加载保存的进度
        self._load_progress()

    def _clean_text(self, text: str) -> str:
        """清理文本中的空白字符"""
        return str(text).strip().replace(" ", "").replace("\t", "").replace("\n", "")
    
    def _save_progress(self):
        """保存进度到文件"""
        try:
            # 确保目录存在
            self.progress_file.parent.mkdir(parents=True, exist_ok=True)
            with open(self.progress_file, 'w', encoding='utf-8') as f:
                json.dump(self.progress, f, ensure_ascii=False, indent=2)
            # 使用debug级别避免过多日志输出
            logger.debug(f"进度已保存: {self.progress.get('current_key', '')} ({self.progress.get('percentage', 0)}%)")
        except Exception as e:
            logger.error(f"保存进度失败: {str(e)}")
    
    def _load_progress(self):
        """从文件加载进度"""
        try:
            if self.progress_file.exists():
                with open(self.progress_file, 'r', encoding='utf-8') as f:
                    saved_progress = json.load(f)
                    self.progress.update(saved_progress)
                    logger.info("已加载保存的进度")
        except Exception as e:
            logger.error(f"加载进度失败: {str(e)}")
    
    def _update_progress(self, **kwargs):
        """更新进度并保存"""
        self.progress.update(kwargs)
        self._save_progress()

    def read_excel(self) -> pd.DataFrame:
        """读取Excel文件并缓存结果"""
        if self.df is None:
            try:
                self.df = pd.read_excel(
                    self.file_name,
                    sheet_name=self.sheet_name,
                    header=0,
                    usecols=self.columns
                )
                # 不使用fillna，我们手动处理合并单元格
                logger.info(f"成功读取Excel文件: {self.file_name}")
            except FileNotFoundError:
                logger.error(f"文件未找到: {self.file_name}")
                raise
            except Exception as e:
                logger.error(f"读取Excel文件时出错: {str(e)}")
                raise
        return self.df

    def check_info(self, json_path: Path):
        """
        校验生成的JSON文件内容
        
        Args:
            json_path: JSON文件路径
            
        Returns:
            problems: 发现的问题数量列表
        """
        with open(json_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        problems = []

        for key, value in data.items():
            # 功能
            functions = []
            # 特性组
            features = []
            for item_index, (item_key, item_value) in enumerate(value.items()):
                if item_index == 0:  # 从第二项开始
                    # 处理角色字符串，统一使用中文逗号分割
                    item_value = item_value.replace(',', '，')
                    role = item_value.strip().split("，")
                    # 去除空字符串并重新检查
                    role = [r.strip() for r in role if r.strip()]
                    if len(role) != 3:
                        problems.append(f"角色信息不全，{key} (实际角色数:{len(role)})")
                        logger.error(f"角色信息不全，{key} (实际角色数:{len(role)}, 角色列表:{role})")
                else:
                    functions.append(item_key)
                    scenario = item_key
                    process = item_value[0]
                    if len(process) != 3:
                        pass
                        # problems += 1
                        problems.append(f"角色信息不全，{key}")
                        logger.error(f"功能过程不是三个，{key}")
        return problems

    def process_data(self):
        try:
            # 重置 DataFrame 缓存，确保读取最新的 Excel 文件
            self.df = None
            
            # 初始化进度
            self._update_progress(
                current_key="正在解析Excel数据",
                total=0,
                current=0,
                percentage=0,
                status="processing"
            )
            
            df = self.read_excel()
            result_dict = OrderedDict()

            # 手动处理合并单元格，前向填充关键字段
            current_func_user_req = None
            current_func_process = None
            current_role = None
            last_valid_role = None  # 记录上一个正常的三个角色
            func_user_req_counter = {}  # 用于跟踪重复的功能用户需求

            # 处理每一行数据
            for idx, row in df.iterrows():
                # 读取当前行的值
                func_user_req = row['功能用户需求'] if pd.notna(row['功能用户需求']) else current_func_user_req
                trigger_event = row['触发事件']
                func_process = row['功能过程'] if pd.notna(row['功能过程']) else current_func_process
                sub_processes = row['子过程描述']
                data_group = row['数据组']
                raw_role = self._clean_text(row['角色']) if pd.notna(row['角色']) else current_role
                
                # 处理角色：确保恰好有3个角色
                if raw_role:
                    role_list = raw_role.replace(',', '，').strip().split("，")
                    role_list = [r.strip() for r in role_list if r.strip()]  # 去除空字符串
                    
                    if len(role_list) > 3:
                        # 如果超过3个角色，使用前3个
                        role_list = role_list[:3]
                        role = "，".join(role_list)
                    elif len(role_list) == 3:
                        # 正好3个角色，记录为有效角色
                        role = "，".join(role_list)
                        last_valid_role = role
                    else:
                        # 不足3个角色，使用上一个有效角色
                        if last_valid_role:
                            role = last_valid_role
                        else:
                            # 如果没有上一个有效角色，补充默认角色
                            while len(role_list) < 3:
                                role_list.append("系统模块")
                            role = "，".join(role_list)
                else:
                    role = current_role

                # 移除重复处理逻辑，保持原始名称
                original_func_user_req = func_user_req

                # 更新当前值（用于下一行的前向填充）
                if pd.notna(row['功能用户需求']):
                    current_func_user_req = func_user_req
                if pd.notna(row['功能过程']):
                    current_func_process = func_process
                if pd.notna(row['角色']):
                    current_role = role


                # 跳过关键字段为空的行
                if not func_user_req or not func_process:
                    logger.warning(f"第{idx+2}行关键字段为空，跳过")
                    continue

                # 检查子过程描述
                if pd.isna(sub_processes) or str(sub_processes).strip() == '':
                    # 如果子过程描述为空，但功能过程存在，我们需要确保至少有一个占位符
                    if func_process and func_user_req:
                        # 初始化功能用户需求字典
                        result_dict.setdefault(func_user_req, OrderedDict())
                        result_dict[func_user_req].setdefault("角色", role)
                        result_dict[func_user_req].setdefault(func_process, [[], []])
                    continue

                # 初始化功能用户需求字典
                result_dict.setdefault(func_user_req, OrderedDict())
                
                # 设置角色
                result_dict[func_user_req].setdefault("角色", role)
                
                # 初始化功能过程
                result_dict[func_user_req].setdefault(func_process, [[], []])
                
                # 添加子过程和数据组
                sub_proc_list, data_group_list = result_dict[func_user_req][func_process]
                # 不进行去重，直接添加子过程描述和数据组
                sub_proc_list.append(sub_processes)
                if data_group:
                    data_group_list.append(data_group)

            # 后处理：确保每个功能过程至少有3个子过程描述
            for func_user_req, req_data in result_dict.items():
                for func_process, process_data in req_data.items():
                    if func_process == "角色":  # 跳过角色字段
                        continue
                    
                    sub_proc_list, data_group_list = process_data
                    
                    # 如果子过程描述少于3个，进行填充
                    while len(sub_proc_list) < 3:
                        if sub_proc_list:
                            # 如果已有描述，重复最后一个
                            sub_proc_list.append(sub_proc_list[-1])
                        else:
                            # 如果没有描述，添加默认描述
                            default_desc = f"执行{func_process}相关操作"
                            sub_proc_list.append(default_desc)

            # 更新进度 - 开始序列化JSON
            self._update_progress(
                current_key="正在序列化JSON数据",
                percentage=90
            )
            
            # 转换为JSON
            json_result = json.dumps(result_dict, ensure_ascii=False, indent=4)
            
            # 更新进度 - 开始保存JSON文件
            self._update_progress(
                current_key="正在保存JSON数据文件",
                percentage=95
            )
          
            with open(self.output_path, 'w', encoding='utf-8') as file:
                file.write(json_result)
                logger.info(f"数据已保存至: {self.output_path}")
            
            # 更新进度为解析完成
            self._update_progress(
                current_key="Excel数据解析完成",
                percentage=100,
                status="completed"
            )
            return json_result

        except Exception as e:
            logger.error(f"处理数据时出错: {str(e)}")
            self._update_progress(
                current_key=f"处理失败: {str(e)}",
                status="error"
            )
            raise
    
    def _get_results_cache_path(self):
        """获取结果缓存文件路径"""
        return Path("web/public/resource/file/output/cache/results.json")
    
    def _load_cached_results(self):
        """加载已缓存的结果"""
        cache_path = self._get_results_cache_path()
        if cache_path.exists():
            try:
                with open(cache_path, 'r', encoding='utf-8') as f:
                    cached_data = json.load(f)
                    logger.info(f"已加载 {len(cached_data)} 个缓存结果")
                    return cached_data
            except Exception as e:
                logger.error(f"加载缓存结果失败: {str(e)}")
        return {}
    
    def _save_cached_results(self, results_dict):
        """保存结果到缓存"""
        cache_path = self._get_results_cache_path()
        try:
            cache_path.parent.mkdir(parents=True, exist_ok=True)
            
            # 记录缓存保存详情
            logger.debug(f"正在保存 {len(results_dict)} 个结果到缓存: {cache_path}")
            
            with open(cache_path, 'w', encoding='utf-8') as f:
                json.dump(results_dict, f, ensure_ascii=False, indent=2, default=str)
            logger.info(f"已保存 {len(results_dict)} 个结果到缓存")
        except Exception as e:
            logger.error(f"保存缓存结果失败: {str(e)}")
    
    def _chapter_to_dict(self, chapter):
        """将Chapter对象转换为字典"""
        return {
            'name': chapter.name,
            'description': chapter.description,
            'functions': chapter.functions,
            'structure': str(chapter.structure),
            'features': [
                {
                    'scenario': f.scenario,
                    'flow_chart_path': str(f.flow_chart),
                    'process': f.process,
                    'input_data': f.input,
                    'output': f.output,
                    'role': f.role
                } for f in chapter.feature
            ]
        }
    
    def _dict_to_chapter(self, chapter_dict):
        """将字典转换为Chapter对象"""
        from .feature import Feature
        features = []
        for f_dict in chapter_dict['features']:
            feature = Feature(
                scenario=f_dict['scenario'],
                flow_chart=f_dict['flow_chart_path'],
                process=f_dict['process'],
                input=f_dict['input_data'],
                output=f_dict['output'],
                role=f_dict['role']
            )
            features.append(feature)
        
        return Chapter(
            name=chapter_dict['name'],
            description=chapter_dict['description'],
            functions=chapter_dict['functions'],
            structure=chapter_dict['structure'],
            feature=features
        )
    
    def clear_cache(self):
        """清理所有缓存，重新开始处理"""
        try:
            # 清理结果缓存
            cache_path = self._get_results_cache_path()
            if cache_path.exists():
                cache_path.unlink()
                logger.info("已清理结果缓存")
            
            # 清理描述缓存
            from .tool import _description_cache_file
            if _description_cache_file.exists():
                _description_cache_file.unlink()
                logger.info("已清理描述缓存")
                
            # 清理图片缓存目录
            img_cache_dir = Path("web/public/resource/file/output/cache/images")
            if img_cache_dir.exists():
                import shutil
                shutil.rmtree(img_cache_dir)
                logger.info("已清理图片缓存")
                
            logger.info("缓存清理完成")
        except Exception as e:
            logger.error(f"清理缓存失败: {str(e)}")
    
    def process_feature(self, item_key, item_value, role):
        """处理单个特性的辅助函数"""
        scenario = item_key
        process = item_value[0]
        flow_chart_path, flow_chart_future = get_flow_chart(role, process)
        input_data = item_value[1][0]
        output = item_value[1][-1]
        return Feature(scenario, flow_chart_path, process, input_data, output, role), flow_chart_future

    def get_data_result(self) -> list:
        """
        高性能并行处理数据封装，支持断点续传
        
        Returns:
            list: 包含 Chapter 对象的列表
        """
        try:
            with open(self.output_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # 加载已缓存的结果
            cached_results = self._load_cached_results()
            result = []
            
            total_items = len(data.items())
            completed_count = len(cached_results)
            
            self._update_progress(
                total=total_items,
                current=completed_count,
                percentage=int((completed_count / total_items) * 100),
                status="processing"
            )
            
            logger.info(f"总共 {total_items} 个项目，已完成 {completed_count} 个，剩余 {total_items - completed_count} 个")
            
            # 先添加已缓存的结果
            for key, cached_chapter in cached_results.items():
                if key in [item[0] for item in data.items()]:
                    chapter = self._dict_to_chapter(cached_chapter)
                    result.append(chapter)
                    logger.info(f"使用缓存结果: {key}")
            
            # 处理未完成的项目
            pending_items = [(key, value) for key, value in data.items() if key not in cached_results]
            
            if not pending_items:
                logger.info("所有项目已完成，直接返回结果")
                self._update_progress(
                    current_key="所有项目已完成",
                    current=total_items,
                    percentage=100,
                    status="completed"
                )
                return result
            
            logger.info(f"开始处理 {len(pending_items)} 个待处理项目")
            
            # 并行处理剩余章节
            with concurrent.futures.ThreadPoolExecutor(max_workers=20) as executor:
                all_futures = {}
                
                for key, value in pending_items:
                    idx = [item[0] for item in data.items()].index(key) + 1
                    logger.info(f"启动处理: {key} ({idx}/{total_items})")
                    
                    functions = []
                    role = None
                    feature_futures = []
                    
                    # 解析数据结构
                    for item_index, (item_key, item_value) in enumerate(value.items()):
                        if item_index == 0:
                            item_value = item_value.replace(',', '，')
                            role = item_value.strip().split("，")
                        else:
                            functions.append(item_key)
                    
                    # 并行启动所有任务
                    description_future = get_description_async(key, functions)
                    structure_path, structure_future = get_structure(key, functions)
                    
                    # 为每个特性启动并行任务
                    for item_index, (item_key, item_value) in enumerate(value.items()):
                        if item_index > 0:  # 跳过角色信息
                            feature_task = executor.submit(
                                self.process_feature,
                                item_key,
                                item_value,
                                role
                            )
                            feature_futures.append(feature_task)
                    
                    # 存储章节的所有 future 任务
                    all_futures[key] = {
                        'idx': idx,
                        'functions': functions,
                        'description_future': description_future,
                        'structure_path': structure_path,
                        'structure_future': structure_future,
                        'feature_futures': feature_futures
                    }
                
                # 收集所有结果并实时保存
                for key, futures_info in all_futures.items():
                    try:
                        idx = futures_info['idx']
                        current_completed = completed_count + len([k for k in all_futures.keys() if k != key and k in cached_results])
                        
                        self._update_progress(
                            current_key=f"处理中: {key}",
                            current=current_completed,
                            percentage=int((current_completed / total_items) * 100),
                            status="processing"
                        )
                        logger.info(f"收集结果: {key} ({idx}/{total_items})")
                        
                        # 等待描述生成完成
                        description = futures_info['description_future'].result()
                        
                        # 等待结构图生成完成  
                        futures_info['structure_future'].result()
                        structure = futures_info['structure_path']
                        
                        # 收集所有特性结果
                        features = []
                        png_futures = []
                        for feature_future in futures_info['feature_futures']:
                            try:
                                feature, png_future = feature_future.result()
                                features.append(feature)
                                png_futures.append(png_future)
                            except Exception as e:
                                logger.error(f"处理特性时出错: {str(e)}")
                                raise
                        
                        # 等待所有图片生成完成
                        for png_future in png_futures:
                            png_future.result()
                        
                        chapter = Chapter(
                            name=key,
                            description=description,
                            functions=futures_info['functions'],
                            structure=structure,
                            feature=features
                        )
                        result.append(chapter)
                        
                        # 实时保存单个结果到缓存
                        cached_results[key] = self._chapter_to_dict(chapter)
                        self._save_cached_results(cached_results)
                        logger.info(f"已保存结果: {key}")
                        
                    except Exception as e:
                        logger.error(f"处理章节 {key} 时出错: {str(e)}")
                        # 不中断整个流程，继续处理其他章节
                        continue
                
            logger.info("数据封装完成")
            self._update_progress(
                current_key="数据封装完成",
                current=total_items,
                percentage=100,
                status="completed"
            )
            return result
            
        except Exception as e:
            logger.error(f"数据封装过程中出错: {str(e)}")
            self._update_progress(
                current_key=f"封装失败: {str(e)}",
                status="error"
            )
            raise

    def get_docx(self, result: list, output_file: str = "web/public/resource/file/output/data.docx", max_chapters: int = None) -> None:
        import time
        start_time = time.time()
        
        try:
            # 如果设置了最大章节数量，则截取结果
            if max_chapters and max_chapters < len(result):
                result = result[:max_chapters]
                logger.info(f"限制章节数量: {max_chapters}")
            
            total_chapters = len(result)
            logger.info(f"开始生成Word文档，章节数量: {total_chapters}")
            
            # 初始化进度
            self._update_progress(
                current_key="准备生成Word文档",
                total=total_chapters,
                current=0,
                percentage=0,
                status="processing"
            )
            
            doc = Document()
            
            # 批量检查所有图片文件是否存在，避免重复文件系统调用
            logger.info("预检查图片文件...")
            image_cache = {}
            unique_paths = set()  # 使用set去重，避免重复检查相同路径
            
            # 收集所有唯一的图片路径
            for item in result:
                structure_path = str(item.structure)
                unique_paths.add(structure_path)
                
                for detail in item.feature:
                    flow_chart_path = str(detail.flow_chart)
                    unique_paths.add(flow_chart_path)
            
            # 批量检查文件存在性
            for path in unique_paths:
                image_cache[path] = Path(path).exists()
                
            logger.info(f"图片文件预检查完成，共检查 {len(image_cache)} 个唯一文件路径")
            
            # 减少进度更新频率 - 每处理10%的章节才更新一次进度
            progress_update_interval = max(1, total_chapters // 10)
            
            # 预定义常用变量减少重复创建
            image_width = Inches(7)
            list_bullet_style = 'List Bullet'
            
            # 预定义常用字符串减少重复创建
            structure_text = '产品结构如图：'
            functions_text = '主要包括如下功能'
            scenario_prefix = '用户场景： '
            flowchart_text = '流程图'
            process_text = '功能过程:'
            input_prefix = '输入：'
            output_prefix = '输出：'
            missing_structure = "[结构图文件不存在]"
            missing_flowchart = "[流程图文件不存在]"
            
            for index, item in enumerate(result, start=1):
                # 只在特定间隔或重要节点更新进度
                if index % progress_update_interval == 0 or index == 1 or index == total_chapters:
                    percentage = int((index - 1) / total_chapters * 95)
                    self._update_progress(
                        current_key=f"生成第{index}/{total_chapters}章",
                        current=index - 1,
                        percentage=percentage
                    )
                
                if index % 100 == 0:  # 每100章输出一次日志
                    elapsed = time.time() - start_time
                    avg_time = elapsed / index
                    remaining_time = avg_time * (total_chapters - index)
                    logger.info(f"已处理 {index}/{total_chapters} 章，平均 {avg_time:.2f}s/章，预计还需 {remaining_time:.1f}s")
                
                # 预处理字符串避免重复格式化
                index_str = str(index)
                item_name = item.name
                
                # 章节标题
                doc.add_heading(f'{index_str}. {item_name}', level=1)
                
                # 1. 产品概述
                doc.add_heading(f'{index_str}.1. 产品概述', level=2)
                doc.add_paragraph(item.description)

                # 2. 产品结构（功能摘要)
                doc.add_heading(f'{index_str}.2. 产品结构（功能摘要)', level=2)
                doc.add_paragraph(structure_text)
                
                # 使用缓存的图片存在状态
                structure_path = str(item.structure)
                if image_cache.get(structure_path, False):
                    doc.add_picture(structure_path, width=image_width)
                else:
                    doc.add_paragraph(missing_structure)
                
                # 功能列表 - 批量添加以提高性能
                doc.add_paragraph(functions_text)
                for fun in item.functions:
                    doc.add_paragraph(fun, style=list_bullet_style)

                # 3. 特性说明
                doc.add_heading(f'{index_str}.3. 特性说明', level=2)

                for loc, detail in enumerate(item.feature, start=1):
                    loc_str = str(loc)
                    scenario = detail.scenario
                    
                    # 特性子标题
                    doc.add_heading(f'{index_str}.3.{loc_str}. {scenario}', level=3)
                    doc.add_paragraph(f'{scenario_prefix}{scenario}', style=list_bullet_style)
                    
                    # 流程图
                    doc.add_paragraph(flowchart_text, style=list_bullet_style)
                    
                    # 使用缓存的图片存在状态
                    flow_chart_path = str(detail.flow_chart)
                    if image_cache.get(flow_chart_path, False):
                        doc.add_picture(flow_chart_path, width=image_width)
                    else:
                        doc.add_paragraph(missing_flowchart)
                    
                    # 功能过程 - 优化字符串处理和循环
                    doc.add_paragraph(process_text, style=list_bullet_style)
                    # 批量处理步骤，减少字符串操作
                    for step_idx, order in enumerate(detail.process, start=1):
                        # 预先处理字符串避免重复操作
                        clean_order = order.rstrip() if order else ''
                        doc.add_paragraph(f'{step_idx}. {clean_order}')
                    
                    # 输入输出 - 使用预定义字符串
                    doc.add_paragraph(f'{input_prefix}{detail.input}', style=list_bullet_style)
                    doc.add_paragraph(f'{output_prefix}{detail.output}', style=list_bullet_style)

            # 更新保存阶段进度
            self._update_progress(
                current_key="正在保存Word文档",
                current=total_chapters,
                percentage=98
            )
            
            # 保存文档
            save_start = time.time()
            doc.save(output_file)
            save_time = time.time() - save_start
            total_time = time.time() - start_time
            
            logger.info(f"Word文档已生成：{output_file}")
            logger.info(f"性能统计 - 总耗时: {total_time:.2f}s, 文档保存: {save_time:.2f}s, 平均: {total_time/total_chapters:.3f}s/章")
            
            # 更新完成进度
            self._update_progress(
                current_key="Word文档生成完成",
                current=total_chapters,
                percentage=100,
                status="completed"
            )
            
        except Exception as e:
            logger.error(f"生成Word文档时出错: {str(e)}")
            logger.error(f"错误详情: {repr(e)}")
            import traceback
            logger.error(f"完整错误栈: {traceback.format_exc()}")
            
            # 更新错误状态
            self._update_progress(
                current_key=f"Word文档生成失败: {str(e)}",
                status="error"
            )
            raise
    
    
    async def get_progress(self):
        """获取当前处理进度"""
        # 确保返回最新的进度
        self._load_progress()
        return self.progress
    
    def reset_progress(self):
        """重置进度"""
        self.progress = {
            "current_key": "",
            "total": 0,
            "current": 0,
            "percentage": 0,
            "status": "idle"
        }
        self._save_progress()
        logger.info("进度已重置")
resolver = MetaDataProcessor(handle_duplicates="merge")  # 合并重复项，不生成版本号